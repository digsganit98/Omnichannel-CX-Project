"""Convert the client-demo overview Markdown into a formatted .docx.

Handles what this specific document actually uses: ATX headings, pipe tables,
fenced code blocks (incl. ASCII diagrams), blockquotes, ordered/unordered lists,
horizontal rules, and inline **bold** / *italic* / `code` / [links](url).
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"
CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
MONO_SIZE = Pt(7.5)


# ── low-level helpers ────────────────────────────────────────────────────────

def shade(cell_or_para, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    target = cell_or_para._tc.get_or_add_tcPr() if hasattr(cell_or_para, "_tc") else cell_or_para._p.get_or_add_pPr()
    target.append(el)


def cell_borders(cell, color="BFBFBF", sz=4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_toc(doc) -> None:
    """Insert a TOC field (levels 1-2). Word populates it on open / F9."""
    head = doc.add_paragraph()
    hr = head.add_run("Contents")
    hr.font.size = Pt(12)
    hr.font.bold = True
    hr.font.color.rgb = ACCENT
    head.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' to build the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)

    brk = doc.add_paragraph()
    brk.paragraph_format.space_after = Pt(0)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    brk.add_run()._r.append(br)


INLINE = re.compile(
    r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str, base_size: Pt | None = None, base_bold=False) -> None:
    """Render inline markdown spans into runs on an existing paragraph."""
    text = text.replace(" ", " ")
    for part in INLINE.split(text):
        if not part:
            continue
        bold, italic, mono, link = base_bold, False, False, False
        body = part
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            bold = italic = True
            body = part[3:-3]
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            bold = True
            body = part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            italic = True
            body = part[1:-1]
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            mono = True
            body = part[1:-1]
        else:
            m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                body, link = m.group(1), True
        run = paragraph.add_run(body)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = CODE_FONT
            run.font.size = base_size or Pt(9)
            run.font.color.rgb = RGBColor(0xA3, 0x15, 0x15)
        else:
            if base_size:
                run.font.size = base_size
            if link:
                run.font.color.rgb = ACCENT
                run.underline = True


# ── block parsing ────────────────────────────────────────────────────────────

def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def main(src: str, out: str) -> None:
    with open(src, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()

    # Base styles
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.06

    for name, size, color, before, after in (
        ("Heading 1", 17, ACCENT, 16, 6),
        ("Heading 2", 13, ACCENT, 13, 5),
        ("Heading 3", 11, RGBColor(0x2E, 0x5C, 0x8A), 10, 4),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.65)

    # Footer page numbers
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run()
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer_p._p.append(fld)

    i, n = 0, len(lines)
    first_h1_done = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── fenced code block ────────────────────────────────────────────
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            while buf and not buf[0].strip():
                buf.pop(0)
            while buf and not buf[-1].strip():
                buf.pop()
            if buf:
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                cell = tbl.rows[0].cells[0]
                shade(cell, CODE_BG)
                cell_borders(cell, color="D9D9D9", sz=2)
                cell.paragraphs[0].text = ""
                for k, cl in enumerate(buf):
                    p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    r = p.add_run(cl.replace("\t", "    "))
                    r.font.name = CODE_FONT
                    r.font.size = MONO_SIZE
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
            continue

        # ── table ────────────────────────────────────────────────────────
        if stripped.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            header = split_row(line)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            ncols = len(header)
            tbl = doc.add_table(rows=1, cols=ncols)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = True
            hdr = tbl.rows[0]
            for c, txt in enumerate(header):
                cell = hdr.cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                add_inline(p, txt, base_size=Pt(8.5), base_bold=True)
                shade(cell, "DCE6F1")
            for r in rows:
                cells = tbl.add_row().cells
                for c in range(ncols):
                    val = r[c] if c < len(r) else ""
                    cell = cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.space_before = Pt(1)
                    add_inline(p, val, base_size=Pt(8.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
            continue

        # ── headings ─────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level, txt = len(m.group(1)), m.group(2).strip()
            if level == 1 and not first_h1_done:
                first_h1_done = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                r.font.size = Pt(22)
                r.font.bold = True
                r.font.color.rgb = ACCENT
                p.paragraph_format.space_after = Pt(10)
                add_toc(doc)
            else:
                h = doc.add_heading(level=min(level, 3))
                h.text = ""
                add_inline(h, txt, base_bold=True)
                keep_with_next(h)
            i += 1
            continue

        # ── horizontal rule ──────────────────────────────────────────────
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "BFBFBF")
            bdr.append(bot)
            pPr.append(bdr)
            i += 1
            continue

        # ── blockquote (may span lines) ──────────────────────────────────
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            text = " ".join(x.strip() for x in buf if x.strip())
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = tbl.rows[0].cells[0]
            shade(cell, "FFF8E1")
            cell_borders(cell, color="E0C97F", sz=4)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_inline(p, text, base_size=Pt(9.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
            continue

        # ── lists ────────────────────────────────────────────────────────
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if m:
            indent, marker, txt = m.group(1), m.group(2), m.group(3)
            ordered = bool(re.fullmatch(r"\d+\.", marker))
            depth = min(len(indent) // 2, 2)
            # gather continuation lines
            i += 1
            while i < n:
                nxt = lines[i]
                if not nxt.strip():
                    break
                if re.match(r"^(\s*)([-*+]|\d+\.)\s+", nxt) or nxt.strip().startswith(("#", "|", ">", "```")):
                    break
                if re.fullmatch(r"-{3,}", nxt.strip()):
                    break
                txt += " " + nxt.strip()
                i += 1
            style = "List Number" if ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22 + 0.2 * depth)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, txt)
            continue

        # ── blank ────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── paragraph (join wrapped lines) ───────────────────────────────
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            if not nxt.strip():
                break
            if nxt.strip().startswith(("#", "|", ">", "```", "- ", "* ", "+ ")):
                break
            if re.match(r"^\s*\d+\.\s+", nxt) or re.fullmatch(r"-{3,}", nxt.strip()):
                break
            buf.append(nxt.strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(buf))

    doc.save(out)
    print(f"saved: {out}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
